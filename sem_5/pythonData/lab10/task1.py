from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_paragraph_font(paragraph, font_name, font_size):
	for run in paragraph.runs:
		run.font.name = font_name
		run.font.size = Pt(font_size)


def main():
	doc = Document()

	p1 = doc.add_paragraph(
		"В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:"
	)
	set_paragraph_font(p1, "Arial", 12)

	p2 = doc.add_paragraph("Флеш-память: используется для хранения скетчей.", style="List Bullet")
	set_paragraph_font(p2, "Arial", 12)

	p3 = doc.add_paragraph("ОЗУ (", style="List Bullet")
	p3.add_run("SRAM").bold = True
	p3.add_run(" — ")
	p3.add_run("static random access memory").italic = True
	p3.add_run(
		", статическая оперативная память с произвольным доступом): используется для хранения и работы переменных.")
	set_paragraph_font(p3, "Arial", 12)

	p4 = doc.add_paragraph("EEPROM (энергозависимая память): используется для хранения постоянной информации.",
						   style="List Bullet")
	set_paragraph_font(p4, "Arial", 12)

	p5 = doc.add_paragraph(
		"Флеш-память и EEPROM являются энергозависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью."
	)
	set_paragraph_font(p5, "Arial", 12)

	table_data = [
		["", "ATmega168", "ATmega328", "ATmega1280", "ATmega2560"],
		["Flash\n(1 кБ flash-памяти занят загрузчиком)", "16 Кбайт", "32 Кбайт", "128 Кбайт", "256 Кбайт"],
		["SRAM", "1 Кбайт", "2 Кбайт", "8 Кбайт", "8 Кбайт"],
		["EEPROM", "512 байт", "1024 байта", "4 Кбайт", "4 Кбайт"],
	]

	table = doc.add_table(rows=0, cols=len(table_data[0]))
	table.style = "Table Grid"

	for row_idx, row_data in enumerate(table_data):
		row = table.add_row()
		for col_idx, cell_data in enumerate(row_data):
			cell = row.cells[col_idx]
			cell.text = cell_data

			for paragraph in cell.paragraphs:
				paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
				if row_idx == 0 or col_idx == 0:
					paragraph.runs[0].bold = True
					shading_elm = OxmlElement('w:shd')
					shading_elm.set(qn('w:fill'), 'D3D3D3')
					cell._element.get_or_add_tcPr().append(shading_elm)

	doc.add_paragraph()

	p6 = doc.add_paragraph()
	p6.add_run(
		"Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25°C. "
		"Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM."
	).italic = True
	set_paragraph_font(p6, "Times New Roman", 12)

	# Сохраняем документ
	output_path = "ATmega_memory.docx"
	doc.save(output_path)


if __name__ == "__main__":
	main()
	print("task1 done")
