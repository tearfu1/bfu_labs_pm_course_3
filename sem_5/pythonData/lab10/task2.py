from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document("ATmega_memory.docx")

    image_paragraph = doc.add_paragraph()
    image_run = image_paragraph.add_run()
    image_run.add_picture("pic.png")
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph()
    caption.add_run("Picture").italic = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in caption.runs:
        run.font.size = Pt(12)

    output_path = "ATmega_memory.docx"
    doc.save(output_path)

if __name__ == "__main__":
    main()
    print("task2 done")
